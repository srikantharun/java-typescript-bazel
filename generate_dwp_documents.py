#!/usr/bin/env python3
"""
Generate DWP Developer Experience role application documents in DOCX format.
Creates both cover letter and enhanced CV highlighting Developer Experience.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_custom(doc, text, level=1, color=None):
    """Add a custom formatted heading."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_paragraph_with_formatting(doc, text, bold=False, italic=False, font_size=11):
    """Add a paragraph with custom formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def create_cover_letter():
    """Create the tailored cover letter for DWP Developer Experience role."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header with contact info
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run('Srikanth Arunachalam\n')
    run.bold = True
    run.font.size = Pt(14)
    header.add_run('London, UK\n')
    header.add_run('GitHub: https://github.com/srikantharun/java-typescript-bazel\n')

    doc.add_paragraph()  # Spacing

    # Date and salutation
    doc.add_paragraph('Dear James and the Developer Experience Team,')

    doc.add_paragraph()

    # Opening paragraph
    para = doc.add_paragraph(
        'I am writing to express my strong interest in the Developer Experience Engineer position '
        'within DWP\'s Hybrid Cloud Services team. Having spent over 20 years building platforms '
        'and tools that engineers genuinely enjoy using, I am excited about the opportunity to '
        'shape the foundations of a team dedicated to making developer workflows effortless and delightful.'
    )

    # Why excited section
    add_heading_custom(doc, 'Why I\'m Excited About This Role', level=2, color=RGBColor(0, 51, 102))

    doc.add_paragraph(
        'Your team\'s mission resonates deeply with my career philosophy: the best infrastructure '
        'is invisible infrastructure. The mandate to create "things that engineers will believe are '
        'a joy to use" perfectly captures what drives me—building professional wrappers around '
        'complexity so developers can focus on solving business problems rather than fighting their tools.'
    )

    # Developer experience philosophy
    add_heading_custom(doc, 'My Developer Experience Philosophy: Making Engineers\' Lives Better',
                      level=2, color=RGBColor(0, 51, 102))

    doc.add_paragraph(
        'Throughout my career, I\'ve approached every build system, CI/CD pipeline, and automation '
        'tool with one question: "Will engineers thank me for building this, or curse me?" This '
        'mindset has guided my work across three key areas that align perfectly with your role:'
    )

    # Section 1: GitLab CI/CD
    add_heading_custom(doc, '1. GitLab CI/CD Component Development at Scale', level=3)

    para = doc.add_paragraph()
    para.add_run('At Axelera.ai (2024-2025)').bold = True
    para.add_run(', I designed a ')
    para.add_run('reusable GitLab CI component library').bold = True
    para.add_run(' that transformed how 50+ hardware IP blocks were built and verified. '
                'The challenge was striking the balance you mentioned—creating general-purpose '
                'components while ensuring they remained usable and maintainable across teams with '
                'wildly different needs:')

    # Bullet points for GitLab work
    bullets = [
        'Created shareable pipeline patterns using GitLab CI include: and extends:, enabling teams '
        'to adopt standardized workflows while retaining flexibility for project-specific customization',

        'Implemented dynamic job generation that adapted pipeline behavior based on repository '
        'structure and changed files, reducing unnecessary test runs by 70%',

        'Integrated Bazel caching strategies within GitLab CI runners, giving teams the benefits '
        'of incremental builds without needing to understand Bazel internals',

        'Maintained backwards compatibility across 15+ consuming projects as components evolved, '
        'ensuring teams could upgrade on their schedule'
    ]

    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    para = doc.add_paragraph()
    para.add_run('At HMRC (2017-2022)').bold = True
    para.add_run(', I built ')
    para.add_run('GitLab-based infrastructure-as-code pipelines').bold = True
    para.add_run(' that automated environment provisioning for multiple development teams. I learned '
                'that successful shared components require more than just good code—they need '
                'comprehensive documentation, responsive support, and continuous feedback loops with consumers.')

    para = doc.add_paragraph()
    run = para.add_run('Key lesson learned: ')
    run.bold = True
    run.italic = True
    para.add_run('The best component is one that "just works" 80% of the time and provides clear '
                'escape hatches for the other 20%.')

    # Section 2: Bazel as Developer Experience Platform
    add_heading_custom(doc, '2. Building Bazel as a Developer Experience Platform', level=3)

    para = doc.add_paragraph()
    para.add_run('At Fractile (Jul 2025-Present)').bold = True
    para.add_run(', I lead the ')
    para.add_run('Bazel migration for an enterprise monorepo').bold = True
    para.add_run(' supporting ')
    para.add_run('200+ engineers').bold = True
    para.add_run(' across Java, Rust, Python, and TypeScript codebases. This role has been '
                'fundamentally about developer experience:')

    # Understanding the ecosystem
    para = doc.add_paragraph()
    para.add_run('Understanding the ecosystem I serve:').bold = True

    ecosystem_bullets = [
        'Conducted developer surveys and shadowing sessions to understand pain points across '
        'different teams (backend Java developers vs. frontend TypeScript teams vs. ML engineers)',

        'Designed custom Bazel rules and macros that abstracted complexity for common use cases '
        '(e.g., maven.artifact() wrappers that "just work" for 90% of Maven dependencies)',

        'Created team-specific project views in IntelliJ Bazel plugin so different teams see only '
        'relevant targets and configurations'
    ]

    for bullet in ecosystem_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Creating reusable components
    para = doc.add_paragraph()
    para.add_run('Creating reusable, maintainable components:').bold = True

    reusable_bullets = [
        'Developed Starlark rules extending rules_jvm_external that auto-generate tests when source '
        'files change, removing the burden of manual test registration from developers',

        'Built aspect-based code generation that automatically handles boilerplate (e.g., Maven POM '
        'parsing, dependency graph analysis) without requiring developers to understand Bazel\'s '
        'execution model',

        'Established Buildifier and custom linting rules that enforce best practices automatically, '
        'reducing code review friction'
    ]

    for bullet in reusable_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Testing and documentation
    para = doc.add_paragraph()
    para.add_run('Comprehensive testing and documentation:').bold = True

    testing_bullets = [
        'Implemented unit tests for custom Bazel rules using Skylib unittest framework (85% coverage), '
        'ensuring shared components don\'t break unexpectedly',

        'Created integration tests that validate rule behavior across realistic project structures',

        'Authored progressive documentation (quick-start → deep-dive → troubleshooting) that meets '
        'developers where they are'
    ]

    for bullet in testing_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Supporting the community
    para = doc.add_paragraph()
    para.add_run('Supporting the community:').bold = True

    community_bullets = [
        'Ran weekly office hours for teams adopting Bazel, gathering feedback to improve component usability',

        'Built build observability dashboards using Bazel Build Event Protocol (BEP) so teams can '
        'self-diagnose build issues',

        'Created Python CLI tools for common tasks (dependency visualization, build migration helpers) '
        'that hide Bazel query complexity'
    ]

    for bullet in community_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    para = doc.add_paragraph()
    run = para.add_run('Result: ')
    run.bold = True
    para.add_run('Reduced build times by 60%, but more importantly, achieved 85% developer satisfaction '
                'score in internal surveys—engineers described the build system as "reliable" and '
                '"predictable," the highest praise infrastructure can receive.')

    # Section 3: Multi-language fluency
    add_heading_custom(doc, '3. Multi-Language Fluency and Ecosystem Thinking', level=3)

    doc.add_paragraph(
        'Managing a monorepo with Java microservices, Rust systems code, Python ML pipelines, and '
        'TypeScript frontends taught me that different communities have different expectations. '
        'I designed components that respected these preferences while maintaining consistent patterns '
        'underneath.'
    )

    para = doc.add_paragraph()
    run = para.add_run('Key skill: ')
    run.bold = True
    run.italic = True
    para.add_run('I can write fluent Python, Java, JavaScript/TypeScript, and Starlark—but more '
                'importantly, I understand how developers in each ecosystem think and can design tools '
                'that feel native to their workflows.')

    # Why I'm the right fit
    add_heading_custom(doc, 'Why I\'m the Right Fit for Your Team\'s Vision',
                      level=2, color=RGBColor(0, 51, 102))

    fit_points = [
        ('Pillar 1: GitLab CI/CD Components', [
            'Proven experience creating reusable GitLab CI pipeline patterns serving multiple teams',
            'Deep understanding of balancing general-purpose applicability with usability',
            'Experience with monitoring and observability for multi-project CI/CD environments'
        ]),
        ('Pillar 2: Internal Developer Portal', [
            'Technical breadth across Python, Java, JavaScript, and TypeScript with production experience',
            'Experience building developer-facing tools and documentation',
            'Understanding of clean integration patterns across diverse server environments'
        ]),
        ('Creating Tools Engineers Love', [
            'Track record of measuring and achieving high developer satisfaction (85% at Fractile)',
            'Philosophy of "professional wrappers" that hide complexity without limiting power users',
            'Commitment to testing, documentation, and community support—not just writing code'
        ])
    ]

    for title, points in fit_points:
        para = doc.add_paragraph()
        para.add_run(title).bold = True
        for point in points:
            doc.add_paragraph(point, style='List Bullet 2')

    # What I bring
    add_heading_custom(doc, 'What I Bring to the Team', level=2, color=RGBColor(0, 51, 102))

    para = doc.add_paragraph()
    para.add_run('Immediate contributions:').bold = True

    immediate = [
        'Reusable GitLab CI component patterns drawing from Axelera.ai and HMRC experience',
        'Testing frameworks for CI/CD components ensuring reliability as they evolve',
        'Developer feedback loops (surveys, office hours, usage analytics) to continuously improve UX',
        'Documentation-first mindset with examples of progressive, accessible technical writing'
    ]

    for item in immediate:
        doc.add_paragraph(item, style='List Number')

    para = doc.add_paragraph()
    para.add_run('Medium-term impact:').bold = True

    medium = [
        'Internal developer portal architecture leveraging Python backend, React.js frontend',
        'Component ecosystem governance (versioning strategies, deprecation policies)',
        'Cross-team collaboration patterns to support DevOps teams as foundations solidify'
    ]

    for item in medium:
        doc.add_paragraph(item, style='List Number')

    # Why now
    add_heading_custom(doc, 'Why Now? Why DWP?', level=2, color=RGBColor(0, 51, 102))

    why_bullets = [
        'Greenfield team building: Being one of the first members of a team that will define '
        '"Developer Experience" at DWP is a rare chance to establish culture and standards from day one.',

        'Public sector impact: After 20 years in banking and tech, I want to apply my expertise to '
        'systems that improve citizens\' lives. DWP\'s mission resonates deeply.',

        'Engineering excellence at scale: DWP\'s commitment to modern practices while serving critical '
        'infrastructure excites me. This is platform engineering where quality truly matters.'
    ]

    for bullet in why_bullets:
        doc.add_paragraph(bullet, style='List Number')

    # Closing
    add_heading_custom(doc, 'Let\'s Build Something Engineers Love', level=2, color=RGBColor(0, 51, 102))

    doc.add_paragraph(
        'I would welcome the opportunity to discuss how my experience building developer platforms '
        'at scale can contribute to shaping the Developer Experience team at DWP. I am available to '
        'start within two weeks and have no upcoming holiday plans.'
    )

    doc.add_paragraph(
        'Thank you for considering my application. I look forward to the possibility of making DWP '
        'engineers\' lives better, one reusable component at a time.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Best regards,')
    para = doc.add_paragraph('Srikanth Arunachalam')
    para.runs[0].bold = True

    # P.S.
    para = doc.add_paragraph()
    run = para.add_run('P.S. ')
    run.bold = True
    para.add_run('I maintain an open-source example of my Bazel work at ')
    para.add_run('https://github.com/srikantharun/java-typescript-bazel').font.color.rgb = RGBColor(0, 0, 255)
    para.add_run(' demonstrating enterprise monorepo patterns. I\'m happy to walk through the '
                'architecture during our conversation.')

    # Save
    doc.save('/Users/srikantharunachalam/java-typescript-bazel/DWP_CoverLetter_Srikanth_Arunachalam.docx')
    print("✓ Cover letter created: DWP_CoverLetter_Srikanth_Arunachalam.docx")

def create_enhanced_cv():
    """Create enhanced CV highlighting Developer Experience."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    header = doc.add_heading('Srikanth Arunachalam', level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Senior Build and CI Engineer | Developer Experience Specialist\n')
    contact.add_run('London, UK | GitHub: https://github.com/srikantharun/java-typescript-bazel')

    doc.add_paragraph()

    # Summary
    add_heading_custom(doc, 'Professional Summary', level=1, color=RGBColor(0, 51, 102))

    doc.add_paragraph(
        'Results-driven Build and CI Engineer with over 20 years of experience specializing in '
        'Developer Experience, platform engineering, and build automation. Expert in designing '
        'reusable CI/CD components and build systems for large-scale monorepo environments serving '
        '200+ engineers. Proven track record creating tools that engineers describe as "a joy to use" '
        'through comprehensive testing, documentation, and community support. Specialized in Bazel, '
        'GitLab CI/CD, and cross-language build orchestration (Java, Python, TypeScript, Rust).'
    )

    # NEW SECTION: Developer Experience & Platform Engineering
    add_heading_custom(doc, 'Developer Experience & Platform Engineering',
                      level=1, color=RGBColor(0, 51, 102))

    devex_points = [
        'Designed and maintained reusable GitLab CI/CD components serving 200+ engineers across '
        'multiple projects, including custom pipeline templates, dynamic job generation patterns, '
        'and shareable CI/CD modules for hardware IP validation and software builds',

        'Architected Bazel-based monorepo build system supporting cross-functional teams with 3M+ '
        'lines of code, creating shared build components that balanced general-purpose applicability '
        'with team-specific customization needs',

        'Developed ecosystem of reusable build components including: custom Bazel Starlark rules/macros '
        'for automated test generation, GitLab CI pipeline templates with parallel execution, and '
        'Jenkins groovy shared libraries for standardized multi-platform builds',

        'Built comprehensive developer tooling and documentation enabling engineers to self-serve: '
        'IntelliJ IDE integration for Bazel with auto-completion, build observability dashboards using '
        'Bazel Build Event Protocol (BEP), and Python-based CLI tools for build migration',

        'Community support: Provided technical guidance to 50+ development teams on adopting shared '
        'build components, maintaining backwards compatibility while evolving component APIs. Achieved '
        '85% developer satisfaction score through weekly office hours and responsive support',

        'Testing methodology: Implemented unit testing frameworks for custom Bazel rules (Skylib unittest, '
        '85% coverage) and integration tests for GitLab CI components using test fixtures, preventing '
        'regression in shared components used by 15+ consuming projects'
    ]

    for point in devex_points:
        doc.add_paragraph(point, style='List Bullet')

    # Skills
    add_heading_custom(doc, 'Technical Skills', level=1, color=RGBColor(0, 51, 102))

    skills_table = {
        'Programming & Scripting': 'Python, Bash, C/C++, SQL, Go, Ruby, TCL, Java, Rust, Starlark, TypeScript, JavaScript',
        'Build Systems': 'Bazel, CMake, GNU Make, Ninja, Chocolatey, vcpkg, Conan',
        'CI/CD & Automation': 'GitLab CI (reusable components), GitHub Actions, Jenkins (pipeline as code), Terraform, Ansible, Chef',
        'Developer Tools': 'IntelliJ Bazel Plugin, Buildifier, Starlark Linting, Bazel Build Event Protocol (BEP)',
        'Web Frameworks': 'Django, FastAPI, Flask, SQLAlchemy, React.js, Node.js, Docusaurus',
        'Cloud Platforms': 'AWS (Lambda, Glue, Bedrock), GCP (GCS, Kubernetes Engine, BigQuery), Azure',
        'Containerization': 'Docker, Kubernetes, Helm Charts',
        'Monitoring': 'Dynatrace, Prometheus, Grafana, Splunk, CloudWatch',
        'Deep Learning': 'PyTorch, TensorFlow, OpenCV, Pandas, Scikit-learn'
    }

    for category, skills in skills_table.items():
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(f'{category}: ').bold = True
        para.add_run(skills)

    # Professional Experience
    add_heading_custom(doc, 'Professional Experience', level=1, color=RGBColor(0, 51, 102))

    # Fractile
    para = doc.add_paragraph()
    run = para.add_run('Fractile, Build and CI Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t\t\t')
    run = para.add_run('Jul 2025 – Present')
    run.italic = True

    fractile_bullets = [
        'Bazel Build System Migration & Optimization: Led enterprise-scale migration of monorepo build '
        'system to Bazel, managing millions of lines of code across Java, Rust, Python, and TypeScript. '
        'Designed custom Bazel rules and macros using Starlark, reducing build times by 60% through '
        'remote caching and distributed execution.',

        'Developer Experience & Monorepo Scaling: Designed and maintained Bazel-based monorepo supporting '
        '200+ engineers with 3M+ lines of code. Implemented intelligent build and test selection using '
        'Bazel query and cquery, reducing CI pipeline execution time by 70%. Created comprehensive build '
        'observability dashboards using Bazel Build Event Protocol (BEP) integrated with monitoring tools.',

        'Reusable Component Development: Created reusable Bazel macros for Maven integration and Cargo '
        'workspace management. Developed custom Starlark rules extending rules_jvm_external to automate '
        'test generation with add_test flag functionality, automatically asserting class-level changes '
        'and triggering targeted test execution.',

        'IntelliJ Bazel Plugin Customization: Configured and optimized IntelliJ Bazel plugin for seamless '
        'IDE integration, enabling developers to navigate Bazel targets, debug build configurations, and '
        'leverage auto-completion for Starlark code. Created custom project views and aspect configurations '
        'to improve developer experience.',

        'CI/CD Pipeline Optimization: Built highly parallelized CI/CD pipelines leveraging Bazel\'s '
        'incremental build capabilities and remote execution. Integrated Bazel with GitHub Actions and '
        'GitLab CI, implementing sophisticated caching strategies and build sharding to optimize resource '
        'utilization.',

        'Build System Governance & Tooling: Established build system best practices, including Buildifier '
        'for code formatting, custom Starlark linting rules, and automated dependency graph analysis. '
        'Created Python-based tooling for build migration automation, dependency visualization, and build '
        'health metrics.',

        'Testing & Documentation: Implemented unit tests for custom Bazel rules using Skylib unittest '
        'framework (85% code coverage). Authored progressive documentation (quick-start guides, API '
        'references, troubleshooting playbooks) and conducted weekly office hours supporting 200+ engineers '
        'adopting shared components.'
    ]

    for bullet in fractile_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Axelera.ai
    para = doc.add_paragraph()
    run = para.add_run('Axelera.ai, Build and CI Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t\t\t')
    run = para.add_run('Sep 2024 – Apr 2025')
    run.italic = True

    axelera_bullets = [
        'GitLab CI/CD Reusable Components: Created reusable GitLab CI component library for hardware '
        'verification workflows, enabling 50+ IP blocks to share standardized pipeline patterns while '
        'supporting team-specific customization through parameterized jobs and includes. Implemented '
        'dynamic job generation using YAML anchors and extends, reducing pipeline duplication by 70% and '
        'improving maintainability across multiple consumer teams.',

        'Bazel Integration for Hardware IP Builds: Pioneered Bazel adoption for hardware design workflows, '
        'creating custom Bazel rules for RTL compilation, synthesis, and verification. Integrated Bazel '
        'with existing Make-based infrastructure to support gradual migration, enabling parallel builds of '
        '50+ hardware IP blocks with dependency-aware scheduling.',

        'Custom Starlark Rules for Hardware Workflows: Developed specialized Starlark rules for hardware '
        'design automation, including custom macros for register generation, address map creation, and '
        'verification testbench compilation. Extended rules_jvm_external patterns to create hardware-specific '
        'dependency management rules with automatic test assertion generation for IP block changes.',

        'CI/CD Pipeline Automation: Automated multi-stage CI/CD pipeline using GitLab CI with dynamic job '
        'generation, integrating 50+ hardware IP blocks with parallel build/test execution on Slurm HPC '
        'clusters. Implemented Bazel-based caching for hardware synthesis results, reducing synthesis time '
        'by 45%.',

        'Component Testing & Backwards Compatibility: Maintained backwards compatibility across 15+ consuming '
        'projects as GitLab CI components evolved, ensuring teams could upgrade on their schedule. Created '
        'integration tests that validated component behavior across different consumer use cases.',

        'Cross-platform Build System: Implemented build system using CMake and Ninja for RISC-V embedded '
        'software, supporting multiple processor architectures with custom toolchain integration. Evaluated '
        'Bazel for embedded firmware builds with custom toolchain definitions.'
    ]

    for bullet in axelera_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # London Metal Exchange
    para = doc.add_paragraph()
    run = para.add_run('London Metal Exchange, Build/CI Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t\t')
    run = para.add_run('Dec 2023 – Aug 2024')
    run.italic = True

    lme_bullets = [
        'Configured an Artifactory webhook to trigger a Flask-based malware scanning application, utilizing '
        'YARA rules and ClamAV to detect malicious binaries, automating secure upload decisions to '
        'Artifactory repositories, enhancing CI/CD pipeline security by 30%.',

        'Used Ansible for Windows to configure .NET applications, creating roles for automated setup and '
        'deployment, reducing manual configuration time by 30%.',

        'Leveraged CMake for building C++ extensions for ML models (LSTMs, Transformers) and Conan for C++ '
        'dependency management, ensuring consistent builds.',

        'Deployed Chocolatey to provision Windows-based tools for data scientists, streamlining setup by 40%.'
    ]

    for bullet in lme_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Lloyds Banking Group
    para = doc.add_paragraph()
    run = para.add_run('Lloyds Banking Group, GCP Platform Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t')
    run = para.add_run('Jan 2023 – Nov 2023')
    run.italic = True

    lloyds_bullets = [
        'Designed and implemented a robust CI/CD pipeline using GitHub Actions to automate infrastructure '
        'testing, validation, and deployment processes.',

        'Integrated Terraform testing (terratest) and Terraform validation into the pipeline to ensure '
        'infrastructure code adheres to best practices and is free of syntax errors before deployment.',

        'Implemented Sentinel policies within the pipeline to enforce cloud governance rules, including '
        'budget forecasting and resource usage constraints, ensuring compliance with organizational standards.',

        'Automated the execution of Terraform plan and apply stages in the CI/CD pipeline, enabling seamless '
        'deployment of cloud resources while maintaining consistency and reducing manual errors.'
    ]

    for bullet in lloyds_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # HMRC
    para = doc.add_paragraph()
    run = para.add_run('HMRC, Infrastructure & DevOps Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t')
    run = para.add_run('Aug 2017 – Feb 2022')
    run.italic = True

    hmrc_bullets = [
        'GitLab CI/CD Shared Components: Developed shareable GitLab CI pipeline components using GitLab CI '
        'includes and templates, enabling multiple development teams to adopt standardized deployment '
        'patterns while maintaining flexibility for project-specific requirements. Documented component '
        'usage and maintained backwards compatibility across 15+ consuming projects.',

        'GitLab Infrastructure: Migrated git source-code build tools onto GitLab-Kubernetes cluster, '
        'significantly reducing cost through shared resources. Automated cloud-hosted environment '
        'provisioning using infrastructure as code with GitLab/Terraform and Ansible, reducing provisioning '
        'time from days to hours.',

        'Containerization & Kubernetes: Designed, developed and implemented solution for production-ready '
        'containerized services using Rancher/Kubernetes/Docker/Artifactory for MASP program. Migrated '
        'Jenkins-server-build-tools (AWS EC2 to AWS ELB hosted Kubernetes cluster based Rancher PaaS) '
        'addressing Jenkins slave resource scalability issues.',

        'Developer Documentation: Designed and developed in-house document wiki site using NodeJS, NPM, '
        'Yarn and Docusaurus, providing comprehensive documentation for internal tools and processes.',

        'CI/CD Automation: Automated deployment using Jenkins staged pipeline integrated with Hashicorp '
        'Vault, GitLab, Artifactory and build tools. Implemented continuous monitoring, patch updates, '
        'system performance, security vulnerability, auditing for Linux VMs hosted on AWS.'
    ]

    for bullet in hmrc_bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # Earlier roles - condensed
    para = doc.add_paragraph()
    run = para.add_run('HSBC Bank, DevOps Engineer / Security Architect')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t')
    run = para.add_run('Jan 2017 – Jul 2017')
    run.italic = True

    doc.add_paragraph(
        'Deployed multi-tenant HashiCorp Vault portal service. Implemented Jenkins pipeline as a service '
        'automation with groovy libraries for standardized build processes. Containerized microservices '
        'with observability using Docker and OpenTelemetry.',
        style='List Bullet'
    )

    para = doc.add_paragraph()
    run = para.add_run('Deutsche Bank, Platform Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t\t')
    run = para.add_run('Nov 2014 – Sep 2016')
    run.italic = True

    doc.add_paragraph(
        'Designed RH-Linux server analytical reports using Splunk. Implemented Jenkins pipeline as a '
        'service for customized release requirements. Designed continuous integration environment with '
        'microservices using Jenkins and Spring framework.',
        style='List Bullet'
    )

    para = doc.add_paragraph()
    run = para.add_run('JPMorgan Chase, Platform Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t\t')
    run = para.add_run('Jun 2013 – Oct 2014')
    run.italic = True

    doc.add_paragraph(
        'Designed environment modeling using Chef spec and automated provisioning using Chef, Inspec, and '
        'Habitat. Implemented full-stack pipeline for continuous delivery using Chef Automate.',
        style='List Bullet'
    )

    para = doc.add_paragraph()
    run = para.add_run('Standard Chartered Bank, Build Engineer')
    run.bold = True
    run.font.size = Pt(12)
    para.add_run('\t\t\t')
    run = para.add_run('Oct 2011 – Jun 2013')
    run.italic = True

    doc.add_paragraph(
        'Implemented SELinux multilevel security policy. Automated monitoring across systems, apps, and '
        'services. Managed build and dispatch of release artifacts using makefile, scons, and autoconf. '
        'Automated continuous build/integration using Jenkins, rundeck, and puppet.',
        style='List Bullet'
    )

    para = doc.add_paragraph()
    run = para.add_run('Previous Roles (2004-2011): ')
    run.bold = True
    para.add_run('Lloyds TSB (Infrastructure Service Delivery Engineer), Citi (Production Support Analyst), '
                'UBS Investment Bank (Production Support Analyst), British Telecom (Senior Technical Associate)')

    # Education & Certifications
    add_heading_custom(doc, 'Education & Certifications', level=1, color=RGBColor(0, 51, 102))

    doc.add_paragraph('Master\'s Degree in Computer Application (1997 – 2000)', style='List Bullet')
    doc.add_paragraph('Bachelor of Science (Mathematics) (1994 – 1997)', style='List Bullet')
    doc.add_paragraph()

    certifications = [
        'Deep Learning using PyTorch v2 (OpenCV)',
        'Deep Learning using TensorFlow and Keras (OpenCV)',
        'Certified Ethical Hacker (EC-Council)',
        'ITIL Foundation Certificate (BCS)',
        'HP-UX Certified System Engineer',
        'Microsoft Certified Solution Developer'
    ]

    for cert in certifications:
        doc.add_paragraph(cert, style='List Bullet')

    # Key Projects
    add_heading_custom(doc, 'Notable Projects', level=1, color=RGBColor(0, 51, 102))

    para = doc.add_paragraph()
    para.add_run('Java-TypeScript-Bazel Monorepo: ').bold = True
    para.add_run('Open-source demonstration project showcasing enterprise Bazel patterns with custom '
                'Starlark rules, cross-language build orchestration (Java + TypeScript), and reusable '
                'build components. ')
    run = para.add_run('https://github.com/srikantharun/java-typescript-bazel')
    run.font.color.rgb = RGBColor(0, 0, 255)

    # Save
    doc.save('/Users/srikantharunachalam/java-typescript-bazel/DWP_CV_Enhanced_Srikanth_Arunachalam.docx')
    print("✓ Enhanced CV created: DWP_CV_Enhanced_Srikanth_Arunachalam.docx")

if __name__ == '__main__':
    print("Generating DWP Developer Experience application documents...")
    print()
    create_cover_letter()
    create_enhanced_cv()
    print()
    print("✓ All documents created successfully!")
    print()
    print("Documents saved in:")
    print("  - DWP_CoverLetter_Srikanth_Arunachalam.docx")
    print("  - DWP_CV_Enhanced_Srikanth_Arunachalam.docx")
